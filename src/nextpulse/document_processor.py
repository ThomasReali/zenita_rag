"""Document processing — multi-format parsing, structural word-chunking, metadata.

Supported content types: PDF, TXT/MD, DOCX, CSV, XLSX, JSON.
- Prose is chunked into a [min, max] window measured by `length_fn` (default = word count;
  the indexer injects a token counter from the embedding model so chunks fit the embedding
  window with NO truncation), cutting at structural boundaries (Art./…, paragraphs) only
  when the chunk is already ≥ min, and force-cutting at the nearest sentence end before max.
- PDFs are parsed per-page so each chunk carries the page where it starts.
- Tables (CSV/XLSX) are chunked by row-groups (header repeated), bounded by word count.
- Per-file errors are caught and reported, never abort the batch (RNF5).
- Light metadata enrichment: doc_type, category, decreto/data (from filename), section.
"""
import io
import json
import logging
import re
from pathlib import Path
from typing import List, Optional, Tuple

from src.nextpulse import config

logger = logging.getLogger("nextpulse.docproc")

SUPPORTED = {".pdf", ".txt", ".md", ".docx", ".csv", ".xlsx", ".json"}
# Files that are metadata *about* the corpus, not content to embed.
SKIP_NAME_SUBSTRINGS = ("manifest", "dettagli_senza_allegati")
# Below this many chars a PDF is treated as a scan/image (needs OCR — out of scope here).
MIN_TEXT_CHARS = 100

# A paragraph starting with one of these markers is a safe place to begin a new chunk.
_LEGAL_BOUNDARY = re.compile(r"^(art\.|articolo|capo|titolo|allegato|sezione)\b", re.I)
_SENTENCE_SPLIT = re.compile(r"(?<=[.!?])\s+")
_DECRETO_RE = re.compile(r"(?:protocollo|numero|n\.?)\s*(\d{1,6})", re.I)
_DATE_RE = re.compile(r"(\d{4})[-_]?(\d{2})[-_]?(\d{2})")


class DocumentProcessor:
    """Load documents of several formats and split them into structural word-chunks."""

    def __init__(self, min_size: int = 200, max_size: int = 480, length_fn=None):
        # Chunk size is measured by length_fn (default: word count). The indexer passes a
        # token counter from the embedding model so chunks never exceed its window.
        self.min_size = min_size
        self.max_size = max_size
        self._len = length_fn or (lambda t: len(t.split()))
        # populated by process_directory()
        self.processed: List[Tuple[str, int]] = []
        self.skipped: List[Tuple[str, str]] = []
        self.failed: List[Tuple[str, str]] = []

    # ── loaders ───────────────────────────────────────────────────────────────

    def load_text(self, file_path: str) -> str:
        with open(file_path, encoding="utf-8-sig", errors="replace") as f:
            return f.read()

    def load_pdf_pages(self, file_path: str) -> List[Tuple[int, str]]:
        """Return [(page_number, text), …] — per page so chunks can cite the page."""
        import pypdf

        reader = pypdf.PdfReader(file_path)
        if reader.is_encrypted:
            try:
                reader.decrypt("")  # many MIT decrees use an empty owner password
            except Exception:
                pass
        pages = [
            (i, page.extract_text() or "")
            for i, page in enumerate(reader.pages, start=1)
        ]
        # Opt-in OCR fallback: fill pages that came back (near) empty — i.e. scans/images.
        if config.OCR_ENABLED:
            pages = self._ocr_fill(file_path, pages)
        return pages

    def _ocr_fill(
        self, file_path: str, pages: List[Tuple[int, str]]
    ) -> List[Tuple[int, str]]:
        """Render the text-poor pages (scanned PDFs) and OCR them with Tesseract.

        Degrades gracefully: if the OCR stack (pymupdf/pytesseract/PIL) or the Tesseract
        binary is unavailable, the original (empty) pages are returned unchanged and the
        document is simply skipped downstream, exactly as before OCR existed (RNF5)."""
        weak = [
            idx for idx, (_, text) in enumerate(pages)
            if len((text or "").strip()) < config.OCR_PAGE_MIN_CHARS
        ]
        if not weak:
            return pages
        try:
            import pytesseract
            from PIL import Image
            try:
                import fitz  # PyMuPDF (preferred import name)
            except ImportError:
                import pymupdf as fitz
        except Exception:
            logger.warning("OCR_ENABLED but the 'ocr' extra is missing "
                            "(uv sync --extra ocr); leaving scanned pages empty")
            return pages
        if config.TESSERACT_CMD:
            pytesseract.pytesseract.tesseract_cmd = config.TESSERACT_CMD
        try:
            doc = fitz.open(file_path)
        except Exception:
            logger.exception("OCR: cannot open %s for rendering", file_path)
            return pages

        out = list(pages)
        for idx in weak:
            page_no = pages[idx][0]
            try:
                pix = doc[page_no - 1].get_pixmap(dpi=config.OCR_DPI)
                img = Image.open(io.BytesIO(pix.tobytes("png")))
                text = pytesseract.image_to_string(img, lang=config.OCR_LANG)
                if text and text.strip():
                    out[idx] = (page_no, text)
            except Exception:
                logger.exception("OCR failed on page %d of %s", page_no, file_path)
        doc.close()
        return out

    def load_pdf(self, file_path: str) -> str:
        return "\n".join(text for _, text in self.load_pdf_pages(file_path))

    def load_docx(self, file_path: str) -> str:
        import docx  # python-docx

        doc = docx.Document(file_path)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(c.text for c in row.cells))
        return "\n".join(parts)

    def load_table(self, file_path: str) -> str:
        """Flat text view of a table (indexing uses _chunk_table for row-grouping)."""
        return "\n\n".join(
            (f"[{sheet}]\n" if sheet else "")
            + "\n".join([" | ".join(map(str, df.columns))]
                        + [" | ".join(map(str, r)) for r in df.values])
            for sheet, df in self._read_table(file_path)
        )

    def load_json(self, file_path: str) -> str:
        with open(file_path, encoding="utf-8") as f:
            data = json.load(f)

        def flatten(obj, prefix="") -> List[str]:
            if isinstance(obj, dict):
                out: List[str] = []
                for k, v in obj.items():
                    out += flatten(v, f"{prefix}{k}.")
                return out
            if isinstance(obj, list):
                out = []
                for i, v in enumerate(obj):
                    out += flatten(v, f"{prefix}{i}.")
                return out
            return [f"{prefix.rstrip('.')}: {obj}"]

        return "\n".join(flatten(data))

    def load_document(self, file_path: str) -> str:
        suffix = Path(file_path).suffix.lower()
        if suffix == ".pdf":
            return self.load_pdf(file_path)
        if suffix in (".txt", ".md"):
            return self.load_text(file_path)
        if suffix == ".docx":
            return self.load_docx(file_path)
        if suffix in (".csv", ".xlsx"):
            return self.load_table(file_path)
        if suffix == ".json":
            return self.load_json(file_path)
        raise ValueError(f"Unsupported format: {suffix}")

    # ── structural word-based chunking ─────────────────────────────────────────

    @staticmethod
    def _paragraphs(text: str) -> List[str]:
        """Group lines into paragraphs; a blank line or a legal marker starts a new one."""
        paras: List[str] = []
        cur: List[str] = []
        for raw in text.split("\n"):
            line = raw.strip()
            if not line:
                if cur:
                    paras.append(" ".join(cur))
                    cur = []
                continue
            if _LEGAL_BOUNDARY.match(line) and cur:
                paras.append(" ".join(cur))
                cur = []
            cur.append(line)
        if cur:
            paras.append(" ".join(cur))
        return [re.sub(r"\s+", " ", p).strip() for p in paras if p.strip()]

    def _make_units(self, text: str, page: Optional[int]) -> List[dict]:
        """Split text into sentence-level units tagged with page/boundary/section."""
        units: List[dict] = []
        for para in self._paragraphs(text):
            is_boundary = bool(_LEGAL_BOUNDARY.match(para))
            section = para[:60].rstrip() if is_boundary else None
            sentences = [s for s in _SENTENCE_SPLIT.split(para) if s.strip()] or [para]
            for i, sentence in enumerate(sentences):
                s = sentence.strip()
                slen = self._len(s)
                parts = [(s, slen)] if slen <= self.max_size else self._split_oversized(s, slen)
                for j, (ptext, plen) in enumerate(parts):
                    first = is_boundary and i == 0 and j == 0
                    units.append({
                        "text": ptext,
                        "len": plen,
                        "page": page,
                        "is_boundary": first,
                        "section": section if first else None,
                    })
        return units

    def _split_oversized(self, text: str, unit_len: int) -> List[Tuple[str, int]]:
        """Hard-split a single oversized unit into pieces that fit max_size (rare path)."""
        words = text.split()
        per = max(1, int(len(words) * self.max_size / unit_len * 0.9))
        return [
            (piece, self._len(piece))
            for i in range(0, len(words), per)
            for piece in [" ".join(words[i:i + per])]
        ]

    def _chunk_units(self, units: List[dict]) -> List[Tuple[str, dict]]:
        """Accumulate units into [min_size, max_size] chunks (boundary + size guardrails).

        - cut at a structural boundary only when the chunk is already ≥ min_size;
        - force a cut (at the current sentence boundary) before exceeding max_size;
        - the tail (or a document shorter than min_size) may stay below min — that's expected.
        """
        chunks: List[Tuple[str, dict]] = []
        buf: List[str] = []
        size = 0
        start: dict = {}

        def flush():
            if buf:
                meta = {}
                if start.get("page") is not None:
                    meta["page"] = start["page"]
                if start.get("section"):
                    meta["section"] = start["section"]
                chunks.append((" ".join(buf), meta))

        for u in units:
            cut_boundary = u["is_boundary"] and size >= self.min_size
            cut_maxsize = size + u["len"] > self.max_size and size > 0  # hard ceiling
            if buf and (cut_boundary or cut_maxsize):
                flush()
                buf, size = [], 0
            if not buf:
                start = {"page": u["page"], "section": u["section"]}
            buf.append(u["text"])
            size += u["len"]
        flush()
        return chunks

    def chunk_text(self, text: str) -> List[str]:
        """Structural word-based chunking of a plain text block (no page metadata)."""
        return [t for t, _ in self._chunk_units(self._make_units(text, None))]

    def process_pdf_bytes(self, raw: bytes, base_meta: dict) -> List[Tuple[str, dict]]:
        """Parse a PDF from raw bytes (no temp file) with the standard token-chunking pipeline.

        Used by the MIT scraper bridge (Approach C) to rechunk downloaded PDF bytes with
        structural boundaries instead of the scraper's character-based chunking.
        `base_meta` provides the RAG-schema fields (source, category, decreto, …);
        `chunk_id` and page/section extras are added per chunk automatically.
        """
        import pypdf

        reader = pypdf.PdfReader(io.BytesIO(raw))
        if reader.is_encrypted:
            try:
                reader.decrypt("")
            except Exception:
                pass
        units: List[dict] = []
        for page_num, page in enumerate(reader.pages, start=1):
            units += self._make_units(page.extract_text() or "", page_num)
        chunks = self._chunk_units(units)
        return [
            (text, {**base_meta, "chunk_id": i, **extra})
            for i, (text, extra) in enumerate(chunks)
        ]

    def process_text_content(self, text: str, base_meta: dict) -> List[Tuple[str, dict]]:
        """Chunk plain text (e.g., extracted from HTML) with the standard pipeline.

        Mirrors `process_pdf_bytes` for non-PDF sources: HTML-extracted text from MIT
        decree pages that have no associated PDF attachment.
        """
        chunks = self._chunk_units(self._make_units(text, None))
        return [
            (txt, {**base_meta, "chunk_id": i, **extra})
            for i, (txt, extra) in enumerate(chunks)
        ]

    # ── tables ──────────────────────────────────────────────────────────────────

    def _read_table(self, file_path: str):
        """Yield (sheet_name, DataFrame) pairs for CSV/XLSX, all cells as strings."""
        import pandas as pd

        suffix = Path(file_path).suffix.lower()
        if suffix == ".csv":
            df = pd.read_csv(
                file_path, sep=None, engine="python", dtype=str,
                keep_default_na=False, encoding="utf-8-sig",
            )
            return [("", df)]
        sheets = pd.read_excel(file_path, sheet_name=None, dtype=str)
        return [(name, df.fillna("")) for name, df in sheets.items()]

    def _chunk_table(self, file_path: str) -> List[str]:
        """Row-group chunks with the header repeated, bounded by the size budget."""
        chunks: List[str] = []
        for sheet, df in self._read_table(file_path):
            header = " | ".join(map(str, df.columns))
            head_size = self._len(header)
            prefix = f"[Foglio: {sheet}]\n" if sheet else ""
            rows: List[str] = []
            size = head_size
            for r in df.values:
                line = " | ".join(map(str, r))
                lw = self._len(line)
                if rows and size + lw > self.max_size:
                    chunks.append(prefix + "\n".join([header] + rows))
                    rows, size = [], head_size
                rows.append(line)
                size += lw
            if rows:
                chunks.append(prefix + "\n".join([header] + rows))
        return chunks

    # ── metadata enrichment ──────────────────────────────────────────────────────

    @staticmethod
    def _doc_metadata(file_path: str) -> dict:
        """Best-effort doc-level metadata from filename/folder (category, decreto, data)."""
        p = Path(file_path)
        name, low, parent = p.name, p.name.lower(), str(p.parent).lower()
        meta: dict = {}
        if "faq" in low:
            meta["category"] = "faq"
        elif "elenco" in low or p.suffix.lower() in (".csv", ".xlsx"):
            meta["category"] = "elenco"
        elif "norm" in low or "norm" in parent or "decreti" in parent \
                or any(k in low for k in ("decreto", "circolare", "legge", "direttiva", "dpcm")):
            meta["category"] = "normativa"
        elif "scheda" in low:
            meta["category"] = "prodotto"
        m = _DECRETO_RE.search(name)
        if m:
            meta["decreto"] = m.group(1)
        d = _DATE_RE.search(name)
        if d:
            meta["data_decreto"] = f"{d.group(1)}-{d.group(2)}-{d.group(3)}"
        return meta

    # ── orchestration ────────────────────────────────────────────────────────────

    def process_document(self, file_path: str) -> List[Tuple[str, dict]]:
        """Parse a single document into (chunk, metadata) tuples."""
        suffix = Path(file_path).suffix.lower()
        base = {
            "source": Path(file_path).name,
            "doc_type": suffix.lstrip("."),
            # Governance: a freshly ingested document is "active" (in vigore). The
            # deterministic audit (scripts/audit_obsolescence.py) flips this to
            # obsolete/poisoned later — never the ingestion, never the LLM.
            "status": "active",
            **self._doc_metadata(file_path),
        }
        # validity_start defaults to the decree date when we could extract it (invariant
        # #4: only real, extracted data — never invented). validity_end stays unset until
        # the audit proves an abrogation.
        if base.get("data_decreto"):
            base.setdefault("validity_start", base["data_decreto"])

        if suffix in (".csv", ".xlsx"):
            chunks = [(t, {}) for t in self._chunk_table(file_path)]
        elif suffix == ".pdf":
            units: List[dict] = []
            for page_num, page_text in self.load_pdf_pages(file_path):
                units += self._make_units(page_text, page_num)
            chunks = self._chunk_units(units)
        else:  # docx / txt / md / json
            chunks = self._chunk_units(self._make_units(self.load_document(file_path), None))

        return [
            (text, {**base, "chunk_id": i, **extra})
            for i, (text, extra) in enumerate(chunks)
        ]

    def candidate_files(self, directory: str) -> List[Path]:
        """Supported, non-manifest files in the tree (file discovery, no parsing)."""
        out: List[Path] = []
        for fp in sorted(Path(directory).rglob("*")):
            if not fp.is_file():
                continue
            if fp.suffix.lower() not in SUPPORTED:
                continue
            if any(s in fp.name.lower() for s in SKIP_NAME_SUBSTRINGS):
                continue
            out.append(fp)
        return out

    def process_directory(self, directory: str) -> List[Tuple[str, dict]]:
        """Process every supported file in a directory tree, robustly (RNF5)."""
        self.processed, self.skipped, self.failed = [], [], []
        all_chunks: List[Tuple[str, dict]] = []

        for file_path in sorted(Path(directory).rglob("*")):
            if not file_path.is_file():
                continue
            name = file_path.name
            suffix = file_path.suffix.lower()

            if suffix not in SUPPORTED:
                self.skipped.append((name, f"estensione non supportata ({suffix or 'n/a'})"))
                continue
            if any(s in name.lower() for s in SKIP_NAME_SUBSTRINGS):
                self.skipped.append((name, "metadati/manifest (riservato a enrichment)"))
                continue

            try:
                chunks = self.process_document(str(file_path))
            except Exception as e:  # noqa: BLE001 — robustness over purity (RNF5)
                self.failed.append((name, f"{type(e).__name__}: {e}"))
                continue

            total_chars = sum(len(c) for c, _ in chunks)
            if not chunks or total_chars == 0:
                self.skipped.append((name, "nessun contenuto estraibile"))
                continue
            if suffix == ".pdf" and total_chars < MIN_TEXT_CHARS:
                self.skipped.append((name, "PDF senza testo (scansione/immagine?)"))
                continue

            all_chunks.extend(chunks)
            self.processed.append((name, len(chunks)))

        return all_chunks
